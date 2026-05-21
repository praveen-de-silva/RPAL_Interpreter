from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.6)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        # header fill
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '404040')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return t

# ════════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CS 3513 – Programming Languages")
run.bold = True
run.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Programming Project: RPAL Interpreter")
run.font.size = Pt(16)

doc.add_paragraph()

# info table on cover
info = doc.add_table(rows=5, cols=2)
info.style = 'Table Grid'
info.alignment = WD_TABLE_ALIGNMENT.CENTER
rows_data = [
    ("Student 1",  "Bula           Index: 230094U"),
    ("Student 2",  "Praveen        Index: 230123K"),
    ("Language",   "C++ (C++17)"),
    ("Executable", "rpal20  /  rpal20.exe"),
    ("Date",       "May 2026"),
]
for i, (label, value) in enumerate(rows_data):
    info.rows[i].cells[0].text = label
    info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    info.rows[i].cells[1].text = value
    for cell in info.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════════
heading("1. Introduction")
para(
    "This report describes the design and implementation of a complete interpreter for the "
    "RPAL (Right-reference Pedagogic Algorithmic Language) programming language. "
    "The interpreter accepts an RPAL source file as input, evaluates it, and produces "
    "output that matches the reference rpal.exe implementation."
)
para(
    "The interpreter is implemented entirely in C++17 without using any lexer/parser "
    "generator tools (no lex, yacc, or equivalent). It is built as a six-stage pipeline:"
)
para("    Lexer  →  Screener  →  Parser  →  Standardizer  →  Flattener  →  CSE Machine",
     italic=True)

heading("Usage", level=2)
code_block("./rpal20 <filename>         # evaluate program")
code_block("./rpal20 -ast <filename>    # print Abstract Syntax Tree and exit")
code_block("./rpal20 -st  <filename>    # print Standardized Tree and exit")

heading("Build", level=2)
code_block("make          # builds the rpal20 executable")
code_block("make clean    # removes the executable")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  2. PROGRAM STRUCTURE
# ════════════════════════════════════════════════════════════════════════════════
heading("2. Program Structure")
para(
    "The interpreter follows a classic linear compiler pipeline. Each stage transforms its "
    "input into a data structure consumed by the next stage."
)

table(
    ["Stage", "Module", "Input", "Output", "Author"],
    [
        ["1", "Lexer",        "Source file (text)",                 "vector<Token>",                    "230094U & 230123K"],
        ["2", "Screener",     "vector<Token> (raw)",                "vector<Token> (filtered)",          "230094U & 230123K"],
        ["3", "Parser",       "vector<Token>",                      "ASTNode* (AST)",                    "230094U & 230123K"],
        ["4", "Standardizer", "ASTNode* (AST)",                     "ASTNode* (ST)",                     "230123K"],
        ["5", "Flattener",    "ASTNode* (ST)",                      "vector<vector<CSENode>>",            "230094U & 230123K"],
        ["6", "CSE Machine",  "vector<vector<CSENode>>",            "Program output (stdout)",           "230094U & 230123K"],
    ],
    col_widths=[1.2, 2.8, 3.5, 4.0, 3.5]
)

heading("Directory Layout", level=2)
code_block(
    "rpal20/\n"
    "├── main.cpp\n"
    "├── Makefile\n"
    "├── lexer/\n"
    "│   ├── Token.h / Token.cpp\n"
    "│   ├── Lexer.h / Lexer.cpp\n"
    "│   └── Screener.h / Screener.cpp\n"
    "├── parser/\n"
    "│   ├── ASTNode.h / ASTNode.cpp\n"
    "│   └── Parser.h / Parser.cpp\n"
    "├── standardizer/\n"
    "│   └── Standardizer.h / Standardizer.cpp\n"
    "├── flattener/\n"
    "│   └── Flattener.h / Flattener.cpp\n"
    "└── cse_machine/\n"
    "    ├── StackValue.h / StackValue.cpp\n"
    "    ├── Environment.h / Environment.cpp\n"
    "    └── CSEMachine.h / CSEMachine.cpp"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  3. MODULE DESCRIPTIONS AND FUNCTION PROTOTYPES
# ════════════════════════════════════════════════════════════════════════════════
heading("3. Module Descriptions and Function Prototypes")

# ── 3.1 Token ──
heading("3.1  Token  (lexer/Token.h, Token.cpp)", level=2)
para("Defines the TokenType enum and the Token data structure used throughout the pipeline.")
code_block(
    "enum class TokenType {\n"
    "    IDENTIFIER, INTEGER, STRING, OPERATOR,\n"
    "    PUNCTUATION, KEYWORD, SPACES, COMMENT, END_OF_FILE\n"
    "};\n\n"
    "struct Token {\n"
    "    TokenType   type;\n"
    "    std::string value;\n"
    "    Token(TokenType t, std::string v);\n"
    "    std::string typeName() const;\n"
    "};"
)

# ── 3.2 Lexer ──
heading("3.2  Lexer  (lexer/Lexer.h, Lexer.cpp)", level=2)
para(
    "Stage 1. Reads the entire source file and converts it into a raw token stream "
    "(including SPACES and COMMENT tokens later removed by the Screener). "
    "Recognises 21 RPAL keywords."
)
code_block(
    "class Lexer {\n"
    "public:\n"
    "    explicit Lexer(const std::string& filename);\n"
    "    std::vector<Token> tokenize();\n"
    "private:\n"
    "    char  currentChar() const;\n"
    "    char  peek()        const;\n"
    "    void  advance();\n"
    "    bool  isLetter       (char c) const;\n"
    "    bool  isDigit        (char c) const;\n"
    "    bool  isOperatorSymbol(char c) const;\n"
    "    Token readIdentifierOrKeyword();\n"
    "    Token readInteger();\n"
    "    Token readString();\n"
    "    Token readOperator();\n"
    "    Token readSpaces();\n"
    "    Token readComment();\n"
    "};"
)

# ── 3.3 Screener ──
heading("3.3  Screener  (lexer/Screener.h, Screener.cpp)", level=2)
para("Stage 2. Filters out SPACES and COMMENT tokens, leaving only meaningful tokens for the parser.")
code_block(
    "class Screener {\n"
    "public:\n"
    "    std::vector<Token> filter(const std::vector<Token>& tokens);\n"
    "};"
)

# ── 3.4 ASTNode ──
heading("3.4  ASTNode  (parser/ASTNode.h, ASTNode.cpp)", level=2)
para(
    "Node type used for both the Abstract Syntax Tree (AST) and the Standardized Tree (ST). "
    "Children are stored as a linked list via child/sibling pointers."
)
code_block(
    "class ASTNode {\n"
    "public:\n"
    "    std::string type;      // e.g. \"let\", \"gamma\", \"IDENTIFIER\"\n"
    "    std::string value;     // leaf value e.g. \"42\", \"x\", \"'hello'\"\n"
    "    ASTNode*    child;\n"
    "    ASTNode*    sibling;\n"
    "    ASTNode(std::string t, std::string v = \"\");\n"
    "    ~ASTNode();\n"
    "    void print(int depth = 0) const;\n"
    "};"
)

# ── 3.5 Parser ──
heading("3.5  Parser  (parser/Parser.h, Parser.cpp)", level=2)
para(
    "Stage 3. Hand-written recursive descent parser implementing the full RPAL grammar. "
    "Uses a tree-building stack; each grammar rule calls buildTree(type, n) to pop n nodes "
    "and attach them as children of a new parent."
)
code_block(
    "class Parser {\n"
    "public:\n"
    "    Parser(const std::vector<Token>& tokenStream);\n"
    "    ASTNode* parse();\n"
    "private:\n"
    "    const Token& nextToken() const;\n"
    "    void read(const std::string& expectedValue = \"\");\n"
    "    void buildTree(const std::string& type, int numChildren);\n"
    "    // 21 recursive descent functions:\n"
    "    void parseE();  void parseEw(); void parseT();  void parseTa();\n"
    "    void parseTc(); void parseB();  void parseBt(); void parseBs();\n"
    "    void parseBp(); void parseA();  void parseAt(); void parseAf();\n"
    "    void parseAp(); void parseR();  void parseRn(); void parseD();\n"
    "    void parseDa(); void parseDr(); void parseDb(); void parseVb();\n"
    "    void parseVl();\n"
    "};"
)

doc.add_page_break()

# ── 3.6 Standardizer ──
heading("3.6  Standardizer  (standardizer/Standardizer.h, Standardizer.cpp)", level=2)
para(
    "Stage 4. Transforms the AST into a Standardized Tree (ST) by applying nine "
    "bottom-up rewrite rules in-place."
)
code_block(
    "class Standardizer {\n"
    "public:\n"
    "    ASTNode* standardize(ASTNode* node);  // recursive entry point\n"
    "private:\n"
    "    ASTNode* makeNode(const std::string& type, const std::string& value = \"\");\n"
    "    ASTNode* copyTree(ASTNode* node);         // deep copy (needed for rec)\n"
    "    ASTNode* getChild(ASTNode* node, int n);  // 0-indexed child access\n"
    "    int      countChildren(ASTNode* node);\n"
    "    void standardizeLet(ASTNode* node);        // Rule 1\n"
    "    void standardizeWhere(ASTNode* node);      // Rule 2\n"
    "    void standardizeFcnForm(ASTNode* node);    // Rule 3\n"
    "    void standardizeLambda(ASTNode* node);     // Rule 4\n"
    "    void standardizeAnd(ASTNode* node);        // Rule 5\n"
    "    void standardizeRec(ASTNode* node);        // Rule 6\n"
    "    void standardizeWithin(ASTNode* node);     // Rule 7\n"
    "    void standardizeAt(ASTNode* node);         // Rule 8\n"
    "    void standardizeEmptyParam(ASTNode* node); // Rule 9\n"
    "};"
)

para("The nine transformation rules:")
rules = [
    ("Rule 1 (let)",      "let x=E in P          →  gamma (lambda x P) E"),
    ("Rule 2 (where)",    "P where x=E           →  gamma (lambda x P) E"),
    ("Rule 3 (fcn_form)", "f x1 x2 = E           →  = f (lambda x1 (lambda x2 E))"),
    ("Rule 4 (lambda)",   "lambda x1 x2 E        →  lambda x1 (lambda x2 E)"),
    ("Rule 5 (and)",      "and (=x1 E1)(=x2 E2)  →  = (tau x1 x2) (tau E1 E2)"),
    ("Rule 6 (rec)",      "rec (= f E)            →  = f (Y* (lambda f E))"),
    ("Rule 7 (within)",   "within (=x1 E1)(=x2 E2) → = x2 (gamma (lambda x1 E2) E1)"),
    ("Rule 8 (@)",        "E @n R                 →  gamma (gamma n E) R"),
    ("Rule 9 (empty)",    "f () = E               →  = f (lambda dummy E)"),
]
for name, rule in rules:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(name + ":  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(rule)
    r2.font.name = "Courier New"
    r2.font.size = Pt(9)

doc.add_paragraph()

# ── 3.7 Flattener ──
heading("3.7  Flattener  (flattener/Flattener.h, Flattener.cpp)", level=2)
para(
    "Stage 5. Converts the ST into delta arrays consumed by the CSE Machine. "
    "Each lambda body and conditional branch becomes its own numbered delta array."
)
code_block(
    "class Flattener {\n"
    "public:\n"
    "    Flattener();\n"
    "    void   flatten(ASTNode* root);\n"
    "    const  std::vector<std::vector<CSENode>>& getDeltas() const;\n"
    "    void   print() const;\n"
    "private:\n"
    "    void flattenNode(ASTNode* node, int currentDelta);\n"
    "    std::vector<std::vector<CSENode>> deltas;\n"
    "    int nextDelta;\n"
    "};"
)

doc.add_page_break()

# ── 3.8 StackValue ──
heading("3.8  StackValue  (cse_machine/StackValue.h, StackValue.cpp)", level=2)
para("Represents any value that can appear on the CSE Machine's value stack.")
code_block(
    "enum class ValueType {\n"
    "    INTEGER, STRING, BOOL, NIL, DUMMY,\n"
    "    TUPLE, CLOSURE, ETA, BUILTIN, PARTIAL\n"
    "};\n\n"
    "class StackValue {\n"
    "public:\n"
    "    static StackValue makeInt    (int v);\n"
    "    static StackValue makeStr    (const std::string& v);\n"
    "    static StackValue makeBool   (bool v);\n"
    "    static StackValue makeNil    ();\n"
    "    static StackValue makeDummy  ();\n"
    "    static StackValue makeTuple  (std::vector<StackValue> elems);\n"
    "    static StackValue makeClosure(int deltaIdx, std::vector<std::string> vars, int envIdx);\n"
    "    static StackValue makeEta    (int deltaIdx, std::vector<std::string> vars, int envIdx);\n"
    "    static StackValue makeBuiltin(const std::string& name);\n"
    "    static StackValue makePartial(const std::string& name, StackValue firstArg);\n"
    "    std::string toString() const;\n"
    "};"
)

# ── 3.9 Environment ──
heading("3.9  Environment  (cse_machine/Environment.h, Environment.cpp)", level=2)
para("Represents a single lexical scope. Environments form a parent chain enabling lexical scoping.")
code_block(
    "class Environment {\n"
    "public:\n"
    "    Environment(int idx, int parentIdx);\n"
    "    void bind      (const std::string& name, const StackValue& val);\n"
    "    bool lookupLocal(const std::string& name, StackValue& out) const;\n"
    "    int  getParentIndex() const;\n"
    "};"
)

# ── 3.10 CSE Machine ──
heading("3.10  CSE Machine  (cse_machine/CSEMachine.h, CSEMachine.cpp)", level=2)
para(
    "Stage 6. Evaluates the delta arrays using the Control–Stack–Environment model. "
    "Implements all 13 CSE rules and 12 built-in functions."
)
code_block(
    "class CSEMachine {\n"
    "public:\n"
    "    explicit CSEMachine(const std::vector<std::vector<CSENode>>& deltaStructures);\n"
    "    void evaluate();\n"
    "private:\n"
    "    // Helpers\n"
    "    int        newEnv  (int parentIdx);\n"
    "    StackValue lookup  (const std::string& name, int envIdx) const;\n"
    "    void       pushDelta(int idx);\n"
    "    // Rules 1,2,3,4,5,8,11,12  (230094U)\n"
    "    void rule1_name      (const CSENode& node);\n"
    "    void rule2_lambda    (const CSENode& node);\n"
    "    void rule3_4_gamma   ();\n"
    "    void rule5_envMarker (const CSENode& node);\n"
    "    void rule8_beta      ();\n"
    "    void rule11_ystar    (StackValue& lambda);\n"
    "    void rule12_eta      (StackValue& eta, StackValue& rand);\n"
    "    // Rules 6,7,9,10,13  (230123K)\n"
    "    void rule6_binaryOp  (const CSENode& node);\n"
    "    void rule7_unaryOp   (const CSENode& node);\n"
    "    void rule9_tau       (const CSENode& node);\n"
    "    void rule10_tupleIndex(StackValue& tuple, StackValue& idx);\n"
    "    void rule13_builtin  (StackValue& rator, StackValue& rand);\n"
    "    // Built-ins  (230123K)\n"
    "    void       builtinPrint       (const StackValue& arg);\n"
    "    StackValue builtinOrder       (const StackValue& arg);\n"
    "    StackValue builtinStem        (const StackValue& arg);\n"
    "    StackValue builtinStern       (const StackValue& arg);\n"
    "    StackValue builtinConc        (const StackValue& s1, const StackValue& s2);\n"
    "    StackValue builtinIsinteger   (const StackValue& arg);\n"
    "    StackValue builtinIsstring    (const StackValue& arg);\n"
    "    StackValue builtinIstruthvalue(const StackValue& arg);\n"
    "    StackValue builtinIstuple     (const StackValue& arg);\n"
    "    StackValue builtinIsfunction  (const StackValue& arg);\n"
    "    StackValue builtinArity       (const StackValue& arg);\n"
    "    StackValue builtinNull        (const StackValue& arg);\n"
    "};"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  4. CSE MACHINE RULES
# ════════════════════════════════════════════════════════════════════════════════
heading("4. CSE Machine – Evaluation Rules")
para(
    "The CSE Machine operates on three data structures: "
    "Control (C) — instruction stack, "
    "Stack (S) — value stack, "
    "Environment (E) — current active environment index."
)

table(
    ["Rule", "Trigger", "Action", "Author"],
    [
        ["1",  "IDENTIFIER on control",        "Lookup in env chain; push value to stack",              "230094U"],
        ["2",  "LAMBDA on control",             "Create closure (delta, vars, env); push to stack",      "230094U"],
        ["3",  "GAMMA + single-param closure",  "New env, bind one param, push delta body",              "230094U"],
        ["4",  "GAMMA + multi-param closure",   "New env, bind params from tuple",                       "230094U"],
        ["5",  "ENV marker on control",         "Restore environment; preserve return value",            "230094U"],
        ["6",  "Binary OPERATOR",               "+  -  *  /  **  gr ge ls le eq ne or & aug",            "230123K"],
        ["7",  "Unary OPERATOR",                "neg / not — pop one value, push result",                "230123K"],
        ["8",  "BETA on control",               "Pop boolean; push correct branch delta",                "230094U"],
        ["9",  "TAU(n) on control",             "Pop n values; create tuple; push to stack",             "230123K"],
        ["10", "GAMMA + tuple on stack",        "Pop integer index; push tuple element (1-based)",       "230123K"],
        ["11", "GAMMA + Y* on stack",           "Wrap lambda in eta (recursive) closure",                "230094U"],
        ["12", "GAMMA + eta closure",           "Unwrap one recursion step; create two new envs",        "230094U"],
        ["13", "GAMMA + built-in on stack",     "Dispatch to built-in function implementation",          "230123K"],
    ],
    col_widths=[1.0, 4.5, 6.5, 2.5]
)

heading("Built-in Functions", level=2)
table(
    ["Function", "Description"],
    [
        ["Print / print",  "Print value to stdout (strings without surrounding quotes)"],
        ["Order",          "Number of elements in a tuple (0 for nil)"],
        ["Stem",           "First character of a string"],
        ["Stern",          "String without its first character"],
        ["Conc",           "Concatenate two strings (curried)"],
        ["Isinteger",      "Type predicate: true if integer"],
        ["Isstring",       "Type predicate: true if string"],
        ["Istruthvalue",   "Type predicate: true if boolean"],
        ["Istuple",        "Type predicate: true if tuple or nil"],
        ["Isfunction",     "Type predicate: true if closure or built-in"],
        ["Arity",          "Number of parameters of a closure"],
        ["null",           "True if nil or empty tuple"],
    ],
    col_widths=[4.0, 11.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  5. SAMPLE TEST CASES
# ════════════════════════════════════════════════════════════════════════════════
heading("5. Sample Test Cases and Output")

tests = [
    ("Comprehensive (factorial, boolean, arithmetic, string)",
     "let\n  logic_test = (true & false) or (not false)\n  and\n  arith_test = 2 ** 3 + 16 / 2 - 1\n  and\n  string_test = 'Hello RPAL'\n  and\n  rec factorial n = n eq 0 -> 1 | n * factorial (n - 1)\nin\n  Print (factorial 5, logic_test, arith_test, string_test)",
     "(120, true, 15, 'Hello RPAL')"),
    ("Infix @ operator",
     "let add x y = x + y\nin Print (3 @add 4)",
     "7"),
    ("Recursive sum",
     "let Sum(A) = Psum (A, Order A)\nwhere rec Psum (T, N) = N eq 0 -> 0\n              | Psum(T, N-1) + T N\nin Print (Sum (1,2,3,4,5))",
     "15"),
    ("String operations (Stem, Stern, Conc)",
     "let s = 'hello'\nin Print (Stem s, Stern s, Conc 'foo' 'bar')",
     "('h', 'ello', 'foobar')"),
    ("Palindrome checker (Stem + Stern + Conc + rec)",
     "let rec Reverse S = S eq '' -> ''\n"
     "                  | Conc (Reverse (Stern S)) (Stem S)\nin\n"
     "let IsPalindrome S = (S eq (Reverse S)) -> 'true' | 'false'\nin\n"
     "Print (IsPalindrome 'madam', IsPalindrome 'hello')",
     "('true', 'false')"),
    ("Tuple operations (aug, nil, Order, null)",
     "let t = nil aug 10 aug 20 aug 30\nin Print (t, Order t, null nil, null t)",
     "((10, 20, 30), 3, true, false)"),
    ("Simultaneous definitions (and) + within",
     "let a = 1 and b = 2 and c = 3 in Print (a+b+c)",
     "6"),
    ("Type predicates",
     "Print (Isinteger 42, Isstring 'hi', Istruthvalue false, Istuple (1,2,3))",
     "(true, true, true, true)"),
]

for title, code, expected in tests:
    heading(title, level=2)
    para("Input:", bold=True)
    code_block(code)
    para("Output:", bold=True)
    code_block(expected)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  6. DIVISION OF RESPONSIBILITIES
# ════════════════════════════════════════════════════════════════════════════════
heading("6. Division of Responsibilities")

table(
    ["Component", "230094U (Bula)", "230123K (Praveen)"],
    [
        ["Lexer / Screener",               "shared", "shared"],
        ["Parser / ASTNode",               "shared", "shared"],
        ["Standardizer (all 9 rules)",     "",       "✓"],
        ["Flattener",                      "shared", "shared"],
        ["CSE Rules 1, 2, 3, 4, 5",       "✓",      ""],
        ["CSE Rules 8, 11, 12 (beta/Y*)",  "✓",      ""],
        ["CSE Rules 6, 7 (operators)",     "",       "✓"],
        ["CSE Rules 9, 10 (tau/tuple)",    "",       "✓"],
        ["CSE Rule 13 + all built-ins",    "",       "✓"],
        ["Environment / StackValue",       "✓ (structure)", "✓ (toString)"],
    ],
    col_widths=[6.5, 4.0, 4.5]
)

# ════════════════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════════════════
out_path = r"C:\Users\ASUS\Desktop\PL_Project\rpal20\report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
